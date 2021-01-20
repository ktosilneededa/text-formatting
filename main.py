import os
import json
from docx import Document
from dictdiffer import diff

filePath = 'sample.docx'
formDataJson = 'formData.json'

sectionOrientationDict = {
    0: "portrait",
    1: "landscape"
}

paragraphAlignmentDict = {
    None: "left",
    0: "center",
    1: "left",
    2: "right",
    3: "justify"
}

lineSpacingDict = {
    None: 1.0,
    0: 1.0,
    1: 1.5,
    2: 2.0,
}


def toCm(value):
    return round(value.cm, 2) if value != 0 else value


def toPt(value):
    return value.pt if value != 0 else value


def zeroIfNone(value):
    return value if value is not None else 0.0


def getBaseStyleProperty(paragraph, property):
    value = getattr(paragraph.paragraph_format, property)
    if value is None:
        value = getattr(paragraph.style.paragraph_format, property)
        if value is None and paragraph.style.base_style is not None:
            value = getattr(paragraph.style.base_style.paragraph_format, property)
            if value is None and paragraph.style.base_style.base_style is None:
                value = None
    return value


def checkFormatting(formdatapath, sampledatapath):
    formData = open(formdatapath)
    sampleData = open(sampledatapath)
    f, s = json.load(formData), json.load(sampleData)
    # os.remove(sampleDataPath)
    print("start\n")
    print("1 / 3: checking page sections formatting...\n")
    d = list(diff(s['sections'], f['sections']))
    if len(d) > 0:
        for i in d:
            print(i)
    else:
        print("ok\n")

    print("\n2 / 3: checking paragraph formatting...")
    g = s['styles']
    e = f['styles']
    pf = []
    for i in e:
        print("\n", i)
        for j in g:
            if j is not None:
                temp2 = i
                temp = next(iter(j.keys()))
                if i == next(iter(j.keys())):
                    pf.append(list(diff(e[i]['paragraphFormat'], j[i]['paragraphFormat'])))
                # else:
                #     print(f"error: style {i} not found")
        if len(pf) == 0:
            print(f"error: style {i} not found")
        else:
            if len(pf[0]) > 0:
                for k in pf[0]:
                    print(k)
            else:
                print("ok")
            pf = []


class Data:
    def __init__(self, file):
        self.document = file
        self.sections = self.document.sections
        self.paragraphs = self.document.paragraphs
        self.text = 0
        self.data = None

    def getSectionProperties(self):
        sectionProperties = [
            {
                "orientation": sectionOrientationDict[self.sections[s].orientation],
                "leftMargin": toCm(self.sections[s].left_margin),
                "rightMargin": toCm(self.sections[s].right_margin),
                "topMargin": toCm(self.sections[s].top_margin),
                "bottomMargin": toCm(self.sections[s].bottom_margin)
            }
            for s in range(len(self.sections))
        ]

        return sectionProperties

    def getParagraphProperties(self):
        paragraphProperties = [
            {
                self.paragraphs[p].style.name: {
                    "paragraph": p + 1,
                    "paragraphFormat": {
                        "alignment": paragraphAlignmentDict[getBaseStyleProperty(self.paragraphs[p], "alignment")],
                        "firstLineIndent": toCm(
                            zeroIfNone(getBaseStyleProperty(self.paragraphs[p], "first_line_indent"))),
                        "lineSpacing": lineSpacingDict[getBaseStyleProperty(self.paragraphs[p], "line_spacing")],
                        "spaceAfter": toPt(zeroIfNone(getBaseStyleProperty(self.paragraphs[p], "space_after"))),
                        "spaceBefore": toPt(zeroIfNone(getBaseStyleProperty(self.paragraphs[p], "space_before"))),
                    }
                }
            }
            if self.paragraphs[p].text != '' else None for p in range(len(self.paragraphs))
        ]

        return paragraphProperties

    def collectData(self):
        self.data = {
            "sections": self.getSectionProperties(),
            "styles": self.getParagraphProperties()
        }

    def makeJsonFile(self):
        self.collectData()
        file = open('sampleData.json', 'w')
        json.dump(self.data, file, indent=4)
        file.close()
        return file.name


def app():
    document = Document(filePath)
    sampleDataJson = Data(document).makeJsonFile()
    checkFormatting(formDataJson, sampleDataJson)


if __name__ == '__main__':
    app()
